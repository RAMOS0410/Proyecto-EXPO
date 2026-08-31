import hashlib
import os
import io
import json
import base64
import sqlite3
import secrets
import pandas as pd
from PIL import Image
from openai import OpenAI
import streamlit as st
import streamlit.components.v1 as components
import docx

# --- CONFIGURACIÓN E ICONO ---
try:
    favicon_img = Image.open("logo.png")
except Exception:
    favicon_img = "🌿"

st.set_page_config(
    page_title="AGRO IA",
    page_icon=favicon_img,
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- ESTILOS CSS ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@500;600;700;800&family=Inter:wght@400;500;600;700&display=swap');

    :root {
        --bg-main: #F4F9F4;
        --card-bg: #FFFFFF;
        --card-border: #D8F3DC;
        --text-title: #081C15;
        --text-body: #1B4332;
        --sidebar-bg: #1B4332;
        --sidebar-bg-2: #2D6A4F;
        --primary-btn: #2D6A4F;
        --primary-btn-hover: #40916C;
    }

    html, body, [data-testid="stAppViewContainer"], .stApp {
        background-color: var(--bg-main) !important;
        color: var(--text-body) !important;
        font-family: 'Inter', sans-serif !important;
    }

    [data-testid="stHeader"] {
        background-color: transparent !important;
    }

    .stApp h1, .stApp h2, .stApp h3, .stApp h4, .stApp h5, .stApp h6,
    [data-testid="stMarkdownContainer"] h1, 
    [data-testid="stMarkdownContainer"] h2, 
    [data-testid="stMarkdownContainer"] h3 {
        color: var(--text-title) !important;
        font-family: 'Poppins', sans-serif !important;
        font-weight: 700 !important;
    }

    .stApp p, .stApp span, .stApp label, .stApp li, [data-testid="stMarkdownContainer"] p, [data-testid="stMarkdownContainer"] li {
        color: #1B4332 !important;
    }

    [data-testid="stChatMessage"] {
        background-color: #FFFFFF !important;
        border: 1px solid var(--card-border) !important;
        border-radius: 12px !important;
        padding: 12px !important;
        margin-bottom: 10px !important;
    }

    [data-testid="stChatMessage"] p, 
    [data-testid="stChatMessage"] li, 
    [data-testid="stChatMessage"] span, 
    [data-testid="stChatMessage"] div {
        color: #081C15 !important;
    }

    input[type="text"], input[type="password"] {
        background-color: #FFFFFF !important;
        color: #081C15 !important;
        border: 1.5px solid var(--card-border) !important;
        border-radius: 10px !important;
    }

    [data-testid="stChatInput"],
    [data-testid="stChatInput"] > div,
    [data-testid="stChatInput"] div[data-baseweb="base-input"],
    [data-testid="stChatInput"] div[data-baseweb="input"] {
        background-color: #FFFFFF !important;
        border-color: #2D6A4F !important;
        border-radius: 16px !important;
    }

    [data-testid="stChatInput"] textarea,
    [data-testid="stChatInput"] input {
        color: #081C15 !important;
        -webkit-text-fill-color: #081C15 !important;
        background-color: transparent !important;
        font-weight: 500 !important;
    }

    [data-testid="stChatInput"] textarea::placeholder,
    [data-testid="stChatInput"] input::placeholder {
        color: #555555 !important;
        -webkit-text-fill-color: #555555 !important;
    }

    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, var(--sidebar-bg) 0%, var(--sidebar-bg-2) 100%) !important;
    }

    [data-testid="stSidebar"] h1, 
    [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] h1 {
        color: #FFFFFF !important;
        font-weight: 800 !important;
    }

    [data-testid="stSidebar"] p, 
    [data-testid="stSidebar"] span, 
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] [data-captioncontainer="true"] {
        color: #D8F3DC !important;
        font-weight: 500 !important;
    }

    [data-testid="stSidebar"] hr {
        border-color: rgba(216, 243, 220, 0.3) !important;
    }

    [data-testid="stSidebar"] div[role="radiogroup"] label {
        background-color: #FFFFFF !important;
        border: 1.5px solid #D8F3DC !important;
        border-radius: 12px !important;
        padding: 12px 16px !important;
        margin-bottom: 8px !important;
        box-shadow: 0 2px 6px rgba(0, 0, 0, 0.05) !important;
        transition: all 0.2s ease !important;
    }

    [data-testid="stSidebar"] div[role="radiogroup"] label p {
        color: #1B4332 !important;
        font-weight: 600 !important;
    }

    [data-testid="stSidebar"] div[role="radiogroup"] label:has(input:checked) {
        background-color: #D8F3DC !important;
        border-color: #40916C !important;
    }

    [data-testid="stSidebar"] div[role="radiogroup"] label:has(input:checked) p {
        color: #081C15 !important;
        font-weight: 800 !important;
    }

    [data-testid="stSidebar"] div[role="radiogroup"] label [data-baseweb="radio"] > div:first-child {
        display: none !important;
    }

    div[data-testid="stVerticalBlock"] > div > div[data-testid="stVerticalBlock"] {
        background-color: var(--card-bg) !important;
        border-radius: 16px !important;
        padding: 24px !important;
        border: 1.5px solid var(--card-border) !important;
        box-shadow: 0 4px 16px rgba(45, 106, 79, 0.06) !important;
    }

    /* BOTONES PRINCIPALES */
    div.stButton > button,
    div.stButton > button *,
    div[data-testid="stDownloadButton"] > button,
    div[data-testid="stDownloadButton"] > button * {
        background-color: var(--primary-btn) !important;
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
        border-radius: 12px !important;
        border: none !important;
        font-family: 'Poppins', sans-serif !important;
        font-weight: 600 !important;
        padding: 0.75rem 1.6rem !important;
        font-size: 15px !important;
        transition: all 0.2s ease !important;
    }

    div.stButton > button:hover,
    div.stButton > button:hover *,
    div[data-testid="stDownloadButton"] > button:hover,
    div[data-testid="stDownloadButton"] > button:hover * {
        background-color: var(--primary-btn-hover) !important;
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(45, 106, 79, 0.2) !important;
    }

    [data-testid="stFileUploader"] {
        background-color: #FFFFFF !important;
        border: 1.5px dashed #40916C !important;
        border-radius: 12px !important;
        padding: 10px !important;
    }

    [data-testid="stFileUploader"] * {
        color: #1B4332 !important;
    }

    div[data-testid="stNotification"] {
        background-color: #E8F5E9 !important;
        color: #1B4332 !important;
        border: 1px solid #B7E4C7 !important;
        border-radius: 12px !important;
    }

    div[data-testid="stNotification"] * {
        color: #1B4332 !important;
    }

    @media screen and (max-width: 768px) {
        .stApp {
            padding: 10px !important;
        }
        
        div[data-testid="stVerticalBlock"] > div > div[data-testid="stVerticalBlock"] {
            padding: 16px !important;
            border-radius: 12px !important;
        }

        [data-testid="stSidebar"] div[role="radiogroup"] label {
            padding: 10px 12px !important;
        }

        [data-testid="stSidebar"] div[role="radiogroup"] label p {
            font-size: 14px !important;
        }

        div.stButton > button, div[data-testid="stDownloadButton"] > button {
            width: 100% !important;
            padding: 0.8rem 1rem !important;
        }

        .stApp h1, [data-testid="stMarkdownContainer"] h1 {
            font-size: 24px !important;
        }

        .stApp h2, [data-testid="stMarkdownContainer"] h2 {
            font-size: 20px !important;
        }
    }
    </style>
""", unsafe_allow_html=True)

# --- FUNCIONES PWA Y LOCALSTORAGE ---
def inject_pwa():
    components.html("""
    <script>
    if (!document.querySelector('link[rel="manifest"]')) {
        const linkManifest = document.createElement('link');
        linkManifest.rel = 'manifest';
        linkManifest.href = 'app/static/manifest.json';
        document.head.appendChild(linkManifest);
    }
    if (!document.querySelector('meta[name="theme-color"]')) {
        const metaTheme = document.createElement('meta');
        metaTheme.name = 'theme-color';
        metaTheme.content = '#2D6A4F';
        document.head.appendChild(metaTheme);
    }
    if ('serviceWorker' in navigator) {
        navigator.serviceWorker.register('app/static/sw.js').catch(() => {});
    }
    </script>
    """, height=0, width=0)

def set_local_storage_token(token):
    components.html(f"""
    <script>
    try {{
        localStorage.setItem('agroia_token', '{token}');
    }} catch (e) {{}}
    </script>
    """, height=0, width=0)

def clear_local_storage_token():
    components.html("""
    <script>
    try {
        localStorage.removeItem('agroia_token');
    } catch (e) {}
    </script>
    """, height=0, width=0)

def try_restore_from_local_storage():
    components.html("""
    <script>
    try {
        const token = localStorage.getItem('agroia_token');
        if (token) {
            const topWindow = window.top;
            const url = new URL(topWindow.location.href);
            if (url.searchParams.get('session_token') !== token) {
                url.searchParams.set('session_token', token);
                topWindow.location.replace(url.toString());
            }
        }
    } catch (e) {}
    </script>
    """, height=0, width=0)

# --- CONEXIÓN OPENAI ---
raw_key = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
api_key = str(raw_key).strip().strip('"').strip("'")
client = OpenAI(api_key=api_key) if api_key else None

# --- BASE DE DATOS SQLITE ---
DB_NAME = "agroia_v3.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT UNIQUE NOT NULL,
            correo TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            nombre_completo TEXT
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sesiones (
            token TEXT PRIMARY KEY,
            usuario TEXT,
            fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS historial (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT,
            cultivo TEXT,
            diagnostico TEXT,
            fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS conversaciones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT,
            titulo TEXT,
            fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS mensajes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversacion_id INTEGER,
            rol TEXT,
            contenido TEXT,
            fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (conversacion_id) REFERENCES conversaciones (id)
        )
    ''')
    conn.commit()
    conn.close()

init_db()

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def registrar_usuario(usuario, correo, password, nombre):
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute("INSERT INTO usuarios (usuario, correo, password, nombre_completo) VALUES (?, ?, ?, ?)",
                       (usuario.strip().lower(), correo.strip().lower(), hash_password(password), nombre))
        conn.commit()
        conn.close()
        return True, "Registro exitoso. Inicia sesión."
    except sqlite3.IntegrityError:
        return False, "El usuario o correo ya existe."
    except Exception as e:
        return False, f"Error: {e}"

def autenticar_usuario(identificador, password):
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute("SELECT id, usuario, correo, password, nombre_completo FROM usuarios WHERE usuario = ? OR correo = ?",
                       (identificador.strip().lower(), identificador.strip().lower()))
        user = cursor.fetchone()

        if user and user[3] == hash_password(password):
            token = secrets.token_hex(16)
            cursor.execute("INSERT INTO sesiones (token, usuario) VALUES (?, ?)", (token, user[1]))
            conn.commit()
            conn.close()
            return user, token, "OK"
        conn.close()
        return None, None, "Usuario o contraseña incorrectos."
    except Exception as e:
        return None, None, f"Error: {e}"

def obtener_usuario_por_token(token):
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT u.usuario, u.nombre_completo
            FROM sesiones s
            JOIN usuarios u ON s.usuario = u.usuario
            WHERE s.token = ?
        """, (token,))
        user = cursor.fetchone()
        conn.close()
        return user
    except Exception:
        return None

def cerrar_sesion_db(token):
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM sesiones WHERE token = ?", (token,))
        conn.commit()
        conn.close()
    except Exception:
        pass

def preparar_imagen(image_pil, max_dim=1024):
    imagen = image_pil.convert("RGB")
    imagen.thumbnail((max_dim, max_dim), Image.LANCZOS)
    return imagen

def encode_image_to_base64(image_pil):
    buffered = io.BytesIO()
    image_pil.save(buffered, format="JPEG", quality=85, optimize=True)
    return base64.b64encode(buffered.getvalue()).decode("utf-8")

def generar_documento_word(datos_json):
    doc = docx.Document()
    doc.add_heading("Informe de Diagnóstico Agrícola - AGRO IA", level=1)
    
    doc.add_heading("Planta y Problema:", level=2)
    doc.add_paragraph(datos_json.get("planta_y_problema", "No especificado"))

    doc.add_heading("Nivel de Gravedad:", level=2)
    doc.add_paragraph(datos_json.get("nivel_gravedad", "No especificado"))

    doc.add_heading("Soluciones Recomendadas:", level=2)
    doc.add_paragraph(datos_json.get("soluciones_recomendadas", "No especificado"))

    doc.add_heading("Prevención:", level=2)
    doc.add_paragraph(datos_json.get("prevencion", "No especificado"))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

inject_pwa()

params = st.query_params

if "autenticado" not in st.session_state:
    if "session_token" in params:
        user_info = obtener_usuario_por_token(params["session_token"])
        if user_info:
            st.session_state.autenticado = True
            st.session_state.usuario = user_info[0]
            st.session_state.nombre_completo = user_info[1] or user_info[0]
            st.session_state.token = params["session_token"]
        else:
            st.session_state.autenticado = False
    else:
        st.session_state.autenticado = False

if "messages" not in st.session_state:
    st.session_state.messages = []

# --- PANTALLA DE ACCESO ---
if not st.session_state.autenticado:
    try_restore_from_local_storage()

    col1, col2, col3 = st.columns([1, 1.3, 1])
    with col2:
        st.markdown('<div style="font-size: 56px; text-align: center; margin-bottom: 6px;">🌿</div>', unsafe_allow_html=True)
        st.markdown('<div style="text-align: center; font-size: 34px; font-weight: 800; color: #081C15;">AGRO IA</div>', unsafe_allow_html=True)
        st.markdown('<div style="text-align: center; color: #1B4332; font-size: 15px; margin-bottom: 24px;">Inteligencia Artificial para el cuidado de tus cultivos</div>', unsafe_allow_html=True)

        tab1, tab2 = st.tabs(["Iniciar Sesión", "Registrarse"])
        with tab1:
            user_input = st.text_input("Usuario o Correo", key="l_user")
            pass_input = st.text_input("Contraseña", type="password", key="l_pass")
            if st.button("Iniciar sesión", use_container_width=True, key="btn_login"):
                user, token, msj = autenticar_usuario(user_input, pass_input)
                if user:
                    st.session_state.autenticado = True
                    st.session_state.usuario = user[1]
                    st.session_state.nombre_completo = user[4] or user[1]
                    st.session_state.token = token

                    st.query_params["session_token"] = token
                    set_local_storage_token(token)
                    st.rerun()
                else:
                    st.error(msj)
        with tab2:
            n_name = st.text_input("Nombre Completo", key="r_name")
            n_user = st.text_input("Usuario", key="r_user")
            n_mail = st.text_input("Correo", key="r_mail")
            n_pass = st.text_input("Contraseña", type="password", key="r_pass")
            if st.button("Crear Cuenta", use_container_width=True, key="btn_register"):
                ok, msj = registrar_usuario(n_user, n_mail, n_pass, n_name)
                if ok:
                    st.success(msj)
                else:
                    st.error(msj)

# --- PANTALLAS INTERNAS ---
else:
    st.sidebar.title("AGRO IA 🌿")
    st.sidebar.caption(f"Usuario activo: {st.session_state.usuario}")
    st.sidebar.write("---")

    opcion_mostrada = st.sidebar.radio(
        "Navegación",
        ["🏠  Inicio e Historial", "🐛  Detectar Plaga", "🤖  Asistente Virtual", "👤  Mi Cuenta"],
        label_visibility="collapsed"
    )
    opcion = opcion_mostrada.split("  ", 1)[1]

    if opcion == "Inicio e Historial":
        st.title(f"¡Bienvenido, {st.session_state.nombre_completo}! 🌿")
        st.caption("Resumen y registro de los diagnósticos aplicados a tus cultivos.")
        st.subheader("Historial de Diagnósticos")

        try:
            conn = sqlite3.connect(DB_NAME)
            df = pd.read_sql_query(
                "SELECT fecha AS Fecha, cultivo AS Cultivo, diagnostico AS Resumen FROM historial WHERE usuario = ? ORDER BY fecha DESC",
                conn,
                params=(st.session_state.usuario,)
            )
            conn.close()

            if not df.empty:
                st.dataframe(df, use_container_width=True)
            else:
                st.info("Aún no has realizado diagnósticos. Selecciona 'Detectar Plaga' en el menú lateral para evaluar una muestra.")
        except Exception as e:
            st.error(f"Error al cargar historial: {e}")

    elif opcion == "Detectar Plaga":
        st.title("Nuevo Diagnóstico Agrícola")
        st.caption("Sube una foto clara de la hoja o cultivo afectado.")

        col1, col2 = st.columns([1, 1])

        with col1:
            st.subheader("Captura de Muestra")
            origen = st.radio("Origen de la imagen:", ["Subir Archivo", "Cámara Directa"])

            imagen_file = None
            if origen == "Subir Archivo":
                imagen_file = st.file_uploader("Formatos permitidos: JPG, PNG", type=["jpg", "png", "jpeg"])
            else:
                imagen_file = st.camera_input("Tomar foto")

            img = None
            if imagen_file:
                img = preparar_imagen(Image.open(imagen_file))
                st.image(img, caption="Muestra seleccionada", use_container_width=True)

        with col2:
            st.subheader("Resultado del Análisis")
            if imagen_file and img is not None:
                if st.button("Ejecutar Análisis", use_container_width=True):
                    with st.spinner("Analizando estado de la planta..."):
                        if client:
                            try:
                                base64_image = encode_image_to_base64(img)
                                prompt_analisis = """
                                Analiza minuciosamente la imagen botánica adjunta antes de dar el diagnóstico. 

                                Pasos de análisis visual obligatorio:
                                1. Inspecciona la morfología foliar (forma de la hoja, bordes, venación, si es trifoliada o simple).
                                2. Observa las lesiones (pústulas, manchas, textura, color).
                                3. Determina con precisión la especie botánica y la patología visualizada.

                                Devuelve la respuesta EXCLUSIVAMENTE en formato JSON estructurado con estas claves exactas:
                                {
                                  "planta_y_problema": "Identifica de forma precisa la especie (ej: Hoja de Frijol) y la plaga o hongo exacto (ej: Roya del frijol / Uromyces appendiculatus)",
                                  "nivel_gravedad": "Bajo, Medio o Alto",
                                  "soluciones_recomendadas": "Tratamientos orgánicos y fungicidas o productos específicos para esta afección",
                                  "prevencion": "Medidas de manejo agrícola preventivas para este cultivo específico"
                                }
                                No agregues etiquetas markdown alrededor del JSON ni texto adicional.
                                """

                                response = client.chat.completions.create(
                                    model="gpt-4o-mini",
                                    messages=[{
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": prompt_analisis},
                                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                                        ]
                                    }],
                                    max_tokens=1200
                                )

                                raw_json = response.choices[0].message.content.strip()
                                if raw_json.startswith("```json"):
                                    raw_json = raw_json[7:]
                                if raw_json.endswith("```"):
                                    raw_json = raw_json[:-3]
                                raw_json = raw_json.strip()

                                datos_json = json.loads(raw_json)
                                st.session_state["ultimo_analisis"] = datos_json

                                conn = sqlite3.connect(DB_NAME)
                                cursor = conn.cursor()
                                cursor.execute("INSERT INTO historial (usuario, cultivo, diagnostico) VALUES (?, ?, ?)",
                                               (st.session_state.usuario, datos_json.get("planta_y_problema", "Desconocido"), raw_json[:100] + "..."))
                                conn.commit()
                                conn.close()

                            except Exception as e:
                                st.error(f"Error durante el proceso: {e}")
                        else:
                            st.error("No se ha configurado la clave API de OpenAI.")

            if "ultimo_analisis" in st.session_state:
                res = st.session_state["ultimo_analisis"]
                st.markdown(f"**🌱 Planta y Problema:** {res.get('planta_y_problema')}")
                st.markdown(f"**⚠️ Nivel de Gravedad:** {res.get('nivel_gravedad')}")
                st.markdown(f"**🛠️ Soluciones Recomendadas:** {res.get('soluciones_recomendadas')}")
                st.markdown(f"**🛡️ Prevención:** {res.get('prevencion')}")

                docx_buffer = generar_documento_word(res)
                st.download_button(
                    label="📄 Descargar Informe Word (.docx)",
                    data=docx_buffer,
                    file_name="diagnostico_agricola.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            elif not imagen_file:
                st.info("Carga o toma una fotografía a la izquierda para desplegar aquí el reporte.")

    elif opcion == "Asistente Virtual":
        st.title("Asistente Agrónomo")
        st.caption("Resuelve tus dudas sobre siembras, fertilizantes y plagas.")

        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()

        cursor.execute("SELECT id, titulo FROM conversaciones WHERE usuario = ? ORDER BY id DESC", (st.session_state.usuario,))
        chats_existentes = cursor.fetchall()

        opciones_map = {"+ Nueva Conversación": None}
        for c in chats_existentes:
            opciones_map[f"Chat #{c[0]}: {c[1]}"] = c[0]

        lista_opciones = list(opciones_map.keys())

        index_seleccionado = 0
        if "current_chat_id" in st.session_state:
            for idx, key in enumerate(lista_opciones):
                if opciones_map[key] == st.session_state.current_chat_id:
                    index_seleccionado = idx
                    break

        chat_seleccionado = st.sidebar.selectbox("Historial de Consultas", lista_opciones, index=index_seleccionado)
        selected_id = opciones_map[chat_seleccionado]

        if selected_id is None:
            if "current_chat_id" in st.session_state and "switch_trigger" not in st.session_state:
                del st.session_state.current_chat_id
                st.session_state.messages = []
        else:
            if st.session_state.get("current_chat_id") != selected_id:
                st.session_state.current_chat_id = selected_id
                cursor.execute("SELECT rol, contenido FROM mensajes WHERE conversacion_id = ? ORDER BY id ASC", (selected_id,))
                mensajes_db = cursor.fetchall()
                st.session_state.messages = [{"role": m[0], "content": m[1]} for m in mensajes_db]

        if "switch_trigger" in st.session_state:
            del st.session_state.switch_trigger

        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        if prompt := st.chat_input("Haz tu consulta agronómica aquí..."):
            if "current_chat_id" not in st.session_state:
                titulo_chat = prompt[:30] + "..." if len(prompt) > 30 else prompt
                cursor.execute("INSERT INTO conversaciones (usuario, titulo) VALUES (?, ?)", (st.session_state.usuario, titulo_chat))
                conn.commit()
                st.session_state.current_chat_id = cursor.lastrowid

            st.session_state.switch_trigger = True
            st.session_state.messages.append({"role": "user", "content": prompt})
            cursor.execute("INSERT INTO mensajes (conversacion_id, rol, contenido) VALUES (?, ?, ?)",
                           (st.session_state.current_chat_id, "user", prompt))
            conn.commit()

            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                if client:
                    try:
                        history = [{"role": "system", "content": "Eres un experto agrónomo que responde en español sencillo, directo y profesional."}] + st.session_state.messages
                        res = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=history,
                            max_tokens=1000
                        )
                        text = res.choices[0].message.content
                        st.markdown(text)

                        st.session_state.messages.append({"role": "assistant", "content": text})
                        cursor.execute("INSERT INTO mensajes (conversacion_id, rol, contenido) VALUES (?, ?, ?)",
                                       (st.session_state.current_chat_id, "assistant", text))
                        conn.commit()
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {e}")
                else:
                    st.error("No hay API Key configurada.")

        conn.close()

    elif opcion == "Mi Cuenta":
        st.title("Perfil de Usuario")
        st.caption("Detalles de tu cuenta de AGRO IA.")

        st.write(f"**Nombre:** {st.session_state.nombre_completo}")
        st.write(f"**Usuario:** {st.session_state.usuario}")
        st.write("---")

        col_logout, _ = st.columns([1, 2])
        with col_logout:
            if st.button("Cerrar sesión", use_container_width=True):
                if "token" in st.session_state:
                    cerrar_sesion_db(st.session_state.token)

                clear_local_storage_token()

                st.session_state.autenticado = False
                st.session_state.usuario = ""
                st.session_state.nombre_completo = ""
                st.session_state.messages = []
                if "current_chat_id" in st.session_state:
                    del st.session_state.current_chat_id

                st.query_params.clear()
                st.rerun()